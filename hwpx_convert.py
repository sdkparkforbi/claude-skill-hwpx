#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
HWP/HWPX 포맷 변환 엔진.

두 갈래로 동작한다.
  (A) 한컴 COM 변환  — hwp/hwpx ↔ pdf/docx/hwp/hwpx/txt/html/rtf.
      한컴오피스 자동화(Open→SaveAs). 레이아웃·표·이미지까지 충실히 보존.
      ★ 작업 흐름의 핵심: 받은 .hwp RFP/양식을 .hwpx로 올려 read/edit 가능하게 만든다.
  (B) 순수 파이썬 추출 — hwp/hwpx → markdown/txt. 한컴 없이 동작(hwpx_read 사용).

주의(SKILL.md): docx→hwpx 방향은 한컴 설치본에 따라 Open이 실패할 수 있다.
그 경우 생성(hwpxgen) 경로를 쓴다.

주요 API
  convert(src, dst)                 # 확장자로 형식 자동판별 (COM)
  hwp_to_hwpx(src, dst=None)        # .hwp → .hwpx
  to_pdf(src, dst=None)             # → .pdf
  to_docx(src, dst=None)            # → .docx
  to_markdown(src, dst=None)        # 순수 파이썬, 표는 마크다운 표로
  to_text(src, dst=None)            # 순수 파이썬 평문(교정추적 .hwpx면 변경 적용본)
  accepted_text(src, dst=None)      # 한컴 GetTextFile = 변경 적용(최종)본의 권위 텍스트

CLI
  python hwpx_convert.py 입력.hwp -o 출력.hwpx
  python hwpx_convert.py 입력.hwpx --md -o 출력.md
  python hwpx_convert.py 입력.hwp --pdf
  python hwpx_convert.py 입력.hwp --final-text -o 최종.txt
"""
import sys, os, io

# 확장자 → 한컴 SaveAs/Open 포맷 문자열
# 주의: 한컴 자동화에는 docx SaveAs 필터가 없다(HWP/HWPX/PDF/HTML/TEXT/RTF만).
# 따라서 .docx는 COM이 아니라 python-docx로 생성한다(to_docx 참조).
_FMT = {'.hwp': 'HWP', '.hwpx': 'HWPX', '.pdf': 'PDF',
        '.txt': 'TEXT', '.html': 'HTML', '.htm': 'HTML', '.rtf': 'RTF'}

def _fmt_of(path):
    ext = os.path.splitext(path)[1].lower()
    if ext not in _FMT:
        raise ValueError('지원하지 않는 형식: %s' % ext)
    return _FMT[ext]

# ───────────────────────── (A) 한컴 COM ─────────────────────────
_HWP = None
def _get_hwp():
    """한컴 자동화 객체 1회 생성(보안모듈 등록 포함). bake와 동일 설정."""
    global _HWP
    if _HWP is not None:
        return _HWP
    import winreg, importlib.util, win32com.client
    spec = importlib.util.find_spec('pyhwpx')
    if spec is None:
        sys.exit("[오류] pyhwpx가 없습니다. 'pip install pyhwpx' 후 다시 실행하세요.")
    dll = os.path.join(os.path.dirname(spec.origin), 'FilePathCheckerModule.dll')
    for p in (r'Software\HNC\HwpAutomation\Modules',
              r'Software\Hnc\HwpUserAction\Modules'):
        k = winreg.CreateKey(winreg.HKEY_CURRENT_USER, p)
        winreg.SetValueEx(k, 'FilePathCheckerModule', 0, winreg.REG_SZ, dll)
        winreg.CloseKey(k)
    hwp = win32com.client.dynamic.Dispatch('HWPFrame.HwpObject')
    hwp.RegisterModule('FilePathCheckDLL', 'FilePathCheckerModule')
    hwp.SetMessageBoxMode(0xFFFFFF)
    _HWP = hwp
    return hwp

def convert(src, dst):
    """src를 열어 dst 확장자 형식으로 저장. .docx는 python-docx 경로로 위임."""
    if os.path.splitext(dst)[1].lower() == '.docx':
        return to_docx(src, dst)
    src = os.path.abspath(src); dst = os.path.abspath(dst)
    in_fmt, out_fmt = _fmt_of(src), _fmt_of(dst)
    hwp = _get_hwp()
    if not hwp.Open(src, in_fmt, ''):
        raise RuntimeError('한컴이 파일을 열지 못했습니다: %s (형식 %s)' % (src, in_fmt))
    os.makedirs(os.path.dirname(dst) or '.', exist_ok=True)
    hwp.SaveAs(dst, out_fmt, '')
    hwp.Clear(1)
    return dst

def _default_out(src, ext):
    return os.path.splitext(src)[0] + ext

def hwp_to_hwpx(src, dst=None):
    return convert(src, dst or _default_out(src, '.hwpx'))
def to_pdf(src, dst=None):
    return convert(src, dst or _default_out(src, '.pdf'))

def to_docx(src, dst=None):
    """hwp/hwpx → docx. 한컴 docx 필터가 없으므로 hwpx_read 구조를 python-docx로 재구성.
    텍스트·표는 보존(정확한 글자모양/색은 미보존). 편집 가능한 docx 산출이 목적."""
    from docx import Document
    import hwpx_read
    dst = dst or _default_out(src, '.docx')
    fmt = hwpx_read.detect_format(src)
    out = Document()
    if fmt == 'hwpx':
        for b in hwpx_read.read_hwpx(src)['blocks']:
            if b['type'] == 'para':
                out.add_paragraph(b['text'])
            else:
                rows = b['rows']
                ncol = max((len(r) for r in rows), default=0)
                if not ncol:
                    continue
                tbl = out.add_table(rows=len(rows), cols=ncol)
                tbl.style = 'Table Grid'
                for i, r in enumerate(rows):
                    for j in range(ncol):
                        tbl.cell(i, j).text = r[j] if j < len(r) else ''
    else:
        for p in hwpx_read.read_hwp(src)['paragraphs']:
            out.add_paragraph(p)
    out.save(dst)
    return dst

def close():
    global _HWP
    if _HWP is not None:
        try: _HWP.Quit()
        except Exception: pass
        _HWP = None

# ───────────────────────── (B) 순수 파이썬 추출 ─────────────────────────
def _rows_to_md(rows):
    if not rows:
        return ''
    ncol = max(len(r) for r in rows)
    def cells(r):
        r = [c.replace('\n', '<br>').replace('|', '\\|') for c in r]
        r += [''] * (ncol - len(r))
        return '| ' + ' | '.join(r) + ' |'
    out = [cells(rows[0]), '| ' + ' | '.join(['---'] * ncol) + ' |']
    out += [cells(r) for r in rows[1:]]
    return '\n'.join(out)

def to_markdown(src, dst=None):
    """hwp/hwpx → 마크다운. hwpx는 표를 마크다운 표로, hwp는 평문 단락."""
    import hwpx_read
    fmt = hwpx_read.detect_format(src)
    lines = []
    if fmt == 'hwpx':
        doc = hwpx_read.read_hwpx(src)
        for b in doc['blocks']:
            if b['type'] == 'para':
                lines.append(b['text'])
            else:
                lines.append(_rows_to_md(b['rows']))
            lines.append('')
    else:
        doc = hwpx_read.read_hwp(src)
        for p in doc['paragraphs']:
            lines.append(p); lines.append('')
    md = '\n'.join(lines).strip() + '\n'
    if dst:
        with io.open(dst, 'w', encoding='utf-8') as f:
            f.write(md)
        return dst
    return md

def to_text(src, dst=None):
    import hwpx_read
    txt = hwpx_read.extract_text(src)
    if dst:
        with io.open(dst, 'w', encoding='utf-8') as f:
            f.write(txt)
        return dst
    return txt

def accepted_text(src, dst=None):
    """한컴으로 연 뒤 GetTextFile로 평문 추출 → **변경 적용(최종)본**의 권위 있는 텍스트.

    핵심(실측): 한컴 GetTextFile('TEXT')은 보기 모드와 무관하게 늘 '변경 내용 적용된'
    최종 텍스트를 돌려준다(삭제 텍스트는 빠지고 삽입은 포함). 또한 자동화에서 변경내용
    수락/거부 액션(RevisionApplyAll 등)은 no-op이라 '원본(거부)본'은 이 방식으로 못 만든다.
    → 편집 전 원문이 필요하면 hwpx_read.read_changes(...)['original']을 쓰고,
      교정추적 .hwpx의 final 추출이 의심스러울 때(orphans>0) 이 함수로 교차검증한다.
    .hwp/.hwpx 모두 직접 입력 가능(변환 불필요)."""
    hwp = _get_hwp()
    src = os.path.abspath(src)
    if not hwp.Open(src, _fmt_of(src), ''):
        raise RuntimeError('한컴이 파일을 열지 못했습니다: %s' % src)
    txt = hwp.GetTextFile('TEXT', '') or ''
    hwp.Clear(1)
    if dst:
        with io.open(dst, 'w', encoding='utf-8') as f:
            f.write(txt)
        return dst
    return txt

# ───────────────────────── CLI ─────────────────────────
def _main(argv):
    args = [a for a in argv if not a.startswith('-')]
    if not args:
        sys.stdout.buffer.write((__doc__ or '').encode('utf-8'))  # cp949 콘솔 크래시 방지
        return 1
    src = args[0]
    out = None
    if '-o' in argv:
        out = argv[argv.index('-o') + 1]
    try:
        if '--md' in argv:
            r = to_markdown(src, out or _default_out(src, '.md'))
        elif '--txt' in argv:
            r = to_text(src, out or _default_out(src, '.txt'))
        elif '--final-text' in argv:
            r = accepted_text(src, out or _default_out(src, '.final.txt'))
        elif '--pdf' in argv:
            r = to_pdf(src, out)
        elif '--docx' in argv:
            r = to_docx(src, out)
        elif '--hwpx' in argv:
            r = hwp_to_hwpx(src, out)
        elif out:
            r = convert(src, out)
        else:
            print('출력 형식을 지정하세요: -o 파일.ext 또는 --md/--txt/--pdf/--docx/--hwpx')
            return 1
        print('[OK] %s → %s' % (src, r))
    finally:
        close()
    return 0

if __name__ == '__main__':
    sys.exit(_main(sys.argv[1:]))
